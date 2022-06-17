from gnews import GNews
from newspaper import Article
from docx import Document
import docx
import datetime
from deep_translator import GoogleTranslator

query = "adriano espaillat"
query = query.title()

def find_lead(paragraph_list):
    for idx,paragraph in enumerate(paragraph_list):
        if len(paragraph.split(" "))>15 and paragraph[-1]=='.':
            del paragraph_list[idx]
            return paragraph

def add_hyperlink(paragraph, url, text, color, underline):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

document = Document()
document.add_paragraph(f"CONGRESSMAN {query.upper()} MENTIONS:")

espaillat_past_24h = []
google_news = GNews(period='24h',max_results=50)
google_news.language = 'english'
espaillat_past_24h += google_news.get_news(query)

print_at_end = []

for article in espaillat_past_24h:
    article_obj = Article(article['url'])

    try:
        article_obj.download()
        article_obj.parse()
    except:
        print_at_end.append(f"error: {article['url']}")
        continue

    date_list = article['published date'].split(" ")
    month_str = date_list[2]
    calendar_map = {'Jun':6,'Jul':7}

    print(f"added: {article['url']}")
    p = document.add_paragraph()

    hyperlink = add_hyperlink(p, article['url'], article_obj.title+f" -- {article['publisher']['title']} {calendar_map[month_str]}/{date_list[1]}/22", '0000FF', True)

    article_text_list = article_obj.text.split("\n")

    # add lead
    p = document.add_paragraph(find_lead(article_text_list))

    for paragraph in article_text_list:
        if query.title().split(" ")[1] in paragraph:
            p = document.add_paragraph(paragraph)


espaillat_past_24h_spanish = []
google_news.language = 'spanish'
espaillat_past_24h_spanish += google_news.get_news(query)

for article in espaillat_past_24h_spanish:
    article_obj = Article(article['url'])

    try:
        article_obj.download()
        article_obj.parse()
    except:
        print_at_end.append(f"error: {article['url']}")
        continue

    date_list = article['published date'].split(" ")
    month_str = date_list[2]
    calendar_map = {'Jun':6,'Jul':7}

    print(article['url'])
    p = document.add_paragraph()
    english_title = GoogleTranslator(source='auto', target='en').translate(article_obj.title)

    hyperlink = add_hyperlink(p, article['url'], english_title+f" -- {article['publisher']['title']} {calendar_map[month_str]}/{date_list[1]}/22", '0000FF', True)

    english_article_text = GoogleTranslator(source='auto', target='en').translate(article_obj.text)
    article_text_list = english_article_text.split("\n")

    # add lead
    p = document.add_paragraph(find_lead(article_text_list))

    for paragraph in article_text_list:
        if query.title().split(" ")[1] in paragraph:
            p = document.add_paragraph(paragraph)

document.add_paragraph()

topics_to_query = ['Housing','Budget','Immigration','Guns']

for topic in topics_to_query:
    google_news = GNews(period='24h',max_results=4)
    google_news.language = 'english'
    document.add_paragraph(f'{topic}:')
    topic_query_result = []
    topic_query_result += google_news.get_news('New York City ' + topic)

    for article in topic_query_result:
        article_obj = Article(article['url'])

        try:
            article_obj.download()
            article_obj.parse()
        except:
            print_at_end.append(f"error: {article['url']}")
            continue
        p = document.add_paragraph()
        hyperlink = add_hyperlink(p, article['url'], article_obj.title+f" -- {article['publisher']['title']} {calendar_map[month_str]}/{date_list[1]}/22", '0000FF', True)

print(print_at_end)

today = datetime.date.today()
document.save(f'Press Clips {today}.docx')