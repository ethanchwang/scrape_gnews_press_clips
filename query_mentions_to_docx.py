from gnews import GNews
from newspaper import Article
from docx import Document
import docx
import datetime
from deep_translator import GoogleTranslator
from dataclasses import dataclass,field
from typing import List

@dataclass
class mention_article:
    url: str=''
    title: str=''
    month_num: int=0
    day_num: int=0
    lead_text: str=''
    publisher: str=''
    mention_text: List = field(default_factory=lambda:[])

@dataclass
class topic_article:
    url: str=''
    title: str=''
    publisher: str=''
    month_num: int=0
    day_num: int=0

def find_lead(paragraph_list):
    for idx,paragraph in enumerate(paragraph_list):
        if len(paragraph.split(" "))>15 and paragraph[-1]=='.':
            del paragraph_list[idx]
            return paragraph

def add_hyperlink(paragraph, url, text, color, underline):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    new_run.bold = True
    new_run.name = 'Arial'
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def query_mentions(query:str, time_frame:str, language:str) -> list:
    all_mentions_articles = []

    #instantiate GNews object
    google_news = GNews(period=time_frame,max_results=10)
    google_news.language = language
    query_results = [] + google_news.get_news(query)

    #add press clips from query_results to document
    for idx,article in enumerate(query_results):
        all_mentions_articles.append(mention_article(url=article['url']))

        article_obj = Article(article['url'])

        #download articles
        try:
            article_obj.download()
            article_obj.parse()
        except:
            print(f"error: {article['url']}")
            continue

        #get date of article
        date_list = article['published date'].split(" ")
        calendar_map = {'Jan':1,'Feb':2,'Mar':3,'Apr':4,'May':5,'Jun':6,'Jul':7,'Aug':8,'Sep':9,'Nov':11,'Dec':12}
        month_num = calendar_map[date_list[2]]
        day_num = date_list[1]

        #ensure clips are in english, if not translate
        article_title = article_obj.title
        if language != 'english':
            article_title = GoogleTranslator(source='auto', target='en').translate(article_obj.title)
            article_text = article_obj.text
            if len(article_obj.text)>5000:
                article_text = article_obj.text[0:4900]
            english_article_text = GoogleTranslator(source='auto', target='en').translate(article_text)
        if language == 'english':
            english_article_text=article_obj.text
        article_text_list = english_article_text.split("\n")

        #populate dataclass
        all_mentions_articles[idx].title = article_title
        all_mentions_articles[idx].month_num = month_num
        all_mentions_articles[idx].day_num = day_num
        all_mentions_articles[idx].publisher = article['publisher']['title']
        all_mentions_articles[idx].lead_text = find_lead(article_text_list)

        #add paragraphs with mention
        for paragraph in article_text_list:
            if query.title().split(" ")[1] in paragraph:
                all_mentions_articles[idx].mention_text.append(paragraph)
    return all_mentions_articles

def query_topics(topics:list,geographical_area:str) -> list:
    all_topic_articles = {}

    for topic in topics:
        all_topic_articles[topic] = []

        google_news = GNews(period='24h',max_results=4)
        google_news.language = 'english'
        topic_query_result = []
        topic_query_result += google_news.get_news(f'{geographical_area} ' + topic)

        for idx,article in enumerate(topic_query_result):
            all_topic_articles[topic].append(topic_article(url=article['url']))
            article_obj = Article(article['url'])

            try:
                article_obj.download()
                article_obj.parse()
            except:
                print(f"error: {article['url']}")
                continue

            #get date of article
            date_list = article['published date'].split(" ")
            calendar_map = {'Jan':1,'Feb':2,'Mar':3,'Apr':4,'May':5,'Jun':6,'Jul':7,'Aug':8,'Sep':9,'Nov':11,'Dec':12}
            month_num = calendar_map[date_list[2]]
            day_num = date_list[1]

            all_topic_articles[topic][idx].title = article_obj.title
            all_topic_articles[topic][idx].month_num = month_num
            all_topic_articles[topic][idx].day_num = day_num
            all_topic_articles[topic][idx].publisher = article['publisher']['title']

    return all_topic_articles

#creates paragraph that bolds name of query
def bold_name(document,name:str,text:str):
    delimiter = name.split()[-1].title()
    split_paragraph = text.split(delimiter)

    p = add_paragraph_without_spacing(document=document)
    for index,text in enumerate(split_paragraph):
        p.add_run(text,style='paragraph_text')
        if index != len(split_paragraph)-1:
            p.add_run(delimiter,style='bold_text')

def add_paragraph_without_spacing(document,text='',style=None):
    if style==None:
        p = document.add_paragraph(text)
    else:
        p = document.add_paragraph(text,style)
    p.paragraph_format.space_before = docx.shared.Pt(0)
    p.paragraph_format.space_after = docx.shared.Pt(0)
    return p

def main():
    document = Document()

    #get name of person to query
    query = "adriano espaillat"
    area = "new york city"
    query = query.title()
    time_frame = "24h"
    topics_to_query = ['Housing','Education','Labor','Budget','Immigration','Guns']

    #instantiate font styles
    style = document.styles
    title_style = style.add_style('title', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    font = title_style.font
    font.size = docx.shared.Pt(11)
    font.name = 'Arial'
    font.bold = True
    font.underline = True

    paragraph_style = style.add_style('paragraph_text', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
    paragraph_font = paragraph_style.font
    paragraph_font.size = docx.shared.Pt(11)
    paragraph_font.name = 'Arial'

    bold_style = style.add_style('bold_text', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
    bold_font = bold_style.font
    bold_font.size = docx.shared.Pt(11)
    bold_font.name = 'Arial'
    bold_font.bold = True

    #title
    p = add_paragraph_without_spacing(document=document,text=f"""CONGRESSMAN {query.split(" ")[1].upper()} MENTIONS:""",style='title')

    #english clips
    for article in query_mentions(query=query,time_frame=time_frame,language='english'):
        p = document.add_paragraph("",style='title')
        p.paragraph_format.space_before = docx.shared.Pt(0)
        p.paragraph_format.space_after = docx.shared.Pt(0)
        add_hyperlink(p, article.url, article.title+f" -- {article.publisher} {article.month_num}/{article.day_num}/22", '0000FF', True)
        p = add_paragraph_without_spacing(document=document)
        p.add_run(article.lead_text,style='paragraph_text')
        p = add_paragraph_without_spacing(document=document)
        for paragraph in article.mention_text:
            bold_name(document=document,name=query,text=paragraph)
            p = add_paragraph_without_spacing(document=document)

    #spanish clips
    for article in query_mentions(query=query,time_frame=time_frame,language='spanish'):
        p = document.add_paragraph("",style='title')
        p.paragraph_format.space_before = docx.shared.Pt(0)
        p.paragraph_format.space_after = docx.shared.Pt(0)
        add_hyperlink(p, article.url, article.title+f" -- {article.publisher} {article.month_num}/{article.day_num}/22", '0000FF', True)
        p = add_paragraph_without_spacing(document=document)
        p.add_run(article.lead_text,style='paragraph_text')
        for paragraph in article.mention_text:
            bold_name(document=document,name=query,text=paragraph)
            p = add_paragraph_without_spacing(document=document)

    #topic clips
    topic_query_results = query_topics(topics=topics_to_query,geographical_area=area)
    p = add_paragraph_without_spacing(document=document,text="PRESS CLIPS:",style='title')
    for topic in topics_to_query:
        p = add_paragraph_without_spacing(document=document)
        p.add_run("-",style='paragraph_text')
        p = add_paragraph_without_spacing(document=document,text=f"{topic.title()}:",style='title')
        for article in topic_query_results[topic]:
            p = document.add_paragraph("",style='title')
            p.paragraph_format.space_before = docx.shared.Pt(0)
            p.paragraph_format.space_after = docx.shared.Pt(0)
            add_hyperlink(p, article.url, article.title+f" -- {article.publisher} {article.month_num}/{article.day_num}/22", '0000FF', True)

    document.save(f'Press Clips {datetime.date.today()}.docx')

if __name__ == '__main__':
    main()